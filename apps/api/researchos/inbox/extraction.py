"""Server-side extraction for Research Inbox uploads."""

from __future__ import annotations

import io
import uuid
from pathlib import Path

import httpx
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from researchos.common.errors import AppError, ValidationError
from researchos.common.secrets import decrypt_secret
from researchos.llm_config.models import LLMProviderConfig

MAX_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_EXTRACTED_CHARS = 100_000
_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".yaml",
    ".yml",
    ".log",
    ".tex",
    ".html",
    ".htm",
}


async def extract_upload(
    db: AsyncSession,
    project_id: uuid.UUID,
    upload: UploadFile,
    *,
    http_client: httpx.AsyncClient | None = None,
) -> tuple[str, str]:
    """Return extracted text and the canonical inbox source type."""

    filename = Path(upload.filename or "upload").name
    suffix = Path(filename).suffix.lower()
    media_type = (upload.content_type or "application/octet-stream").lower()
    data = await upload.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValidationError("Upload exceeds the 25 MB Research Inbox limit.")
    if not data:
        raise ValidationError("The uploaded file is empty.")

    if media_type.startswith("audio/") or suffix in {".mp3", ".wav", ".m4a", ".ogg", ".webm"}:
        text = await transcribe_audio(
            db,
            project_id,
            filename=filename,
            media_type=media_type,
            data=data,
            http_client=http_client,
        )
        return _bounded(text), "audio_transcript"
    if suffix == ".pdf" or media_type == "application/pdf":
        return _bounded(_extract_pdf(data)), "file"
    docx_media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if suffix == ".docx" or media_type == docx_media_type:
        return _bounded(_extract_docx(data)), "file"
    if suffix in _TEXT_EXTENSIONS or media_type.startswith("text/"):
        return _bounded(_decode_text(data)), "file"
    raise ValidationError(
        "Unsupported file type. Upload text, Markdown, CSV, JSON, YAML, TeX, "
        "HTML, DOCX, PDF, or audio."
    )


def _bounded(text: str) -> str:
    normalized = text.replace("\x00", "").strip()
    if not normalized:
        raise ValidationError("No readable text could be extracted from the upload.")
    if len(normalized) > MAX_EXTRACTED_CHARS:
        return normalized[:MAX_EXTRACTED_CHARS]
    return normalized


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValidationError("The text file encoding is not supported.")


def _extract_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"[Page {number}]\n{text}")
        return "\n\n".join(pages)
    except Exception as exc:  # noqa: BLE001 - normalize parser failures
        raise ValidationError(f"The PDF could not be parsed: {str(exc)[:200]}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
        blocks = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                values = [" ".join(cell.text.split()) for cell in row.cells]
                if any(values):
                    blocks.append(" | ".join(values))
        return "\n\n".join(blocks)
    except Exception as exc:  # noqa: BLE001 - normalize parser failures
        raise ValidationError(f"The DOCX could not be parsed: {str(exc)[:200]}") from exc


async def transcribe_audio(
    db: AsyncSession,
    project_id: uuid.UUID,
    *,
    filename: str,
    media_type: str,
    data: bytes,
    http_client: httpx.AsyncClient | None = None,
) -> str:
    """Use the project's explicit ``asr`` OpenAI-compatible provider."""

    config = await db.scalar(
        select(LLMProviderConfig)
        .where(
            LLMProviderConfig.project_id == project_id,
            LLMProviderConfig.name == "asr",
            LLMProviderConfig.is_active.is_(True),
        )
        .order_by(LLMProviderConfig.updated_at.desc())
        .limit(1)
    )
    if config is None:
        raise AppError(
            "Audio transcription is not configured. Add an active OpenAI-compatible "
            "provider named 'asr' in Management Center.",
            code="asr_not_configured",
            http_status=409,
        )
    if config.provider_type != "openai_compatible":
        raise ValidationError("The ASR provider must use the OpenAI-compatible protocol.")
    api_key = decrypt_secret(config.api_key)
    if not api_key:
        raise ValidationError("The ASR provider has no API key.")

    owns_client = http_client is None
    client = http_client or httpx.AsyncClient(timeout=120)
    try:
        response = await client.post(
            f"{config.base_url.rstrip('/')}/audio/transcriptions",
            headers={"Authorization": f"Bearer {api_key}"},
            files={"file": (filename, data, media_type)},
            data={"model": config.model or "whisper-1", "response_format": "json"},
        )
        response.raise_for_status()
        payload = response.json()
        text = payload.get("text") if isinstance(payload, dict) else None
        if not isinstance(text, str) or not text.strip():
            raise AppError(
                "The ASR provider returned no transcript.",
                code="asr_empty_response",
                http_status=502,
            )
        return text
    except httpx.HTTPStatusError as exc:
        raise AppError(
            f"The ASR provider returned HTTP {exc.response.status_code}.",
            code="asr_provider_error",
            http_status=502,
        ) from exc
    except httpx.TimeoutException as exc:
        raise AppError(
            "Audio transcription timed out.",
            code="asr_timeout",
            http_status=504,
        ) from exc
    except httpx.RequestError as exc:
        raise AppError(
            "Could not connect to the ASR provider.",
            code="asr_unavailable",
            http_status=502,
        ) from exc
    finally:
        if owns_client:
            await client.aclose()
